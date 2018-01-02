import docx
import os


def getText(name):
    qName = r"D:\\OUT\\"+name+".docx"
    aName = r"D:\\OUT\\"+name+"答案.doc"
    try:
        file = docx.Document(qName)
        print("段落数:" + str(len(file.paragraphs)))
        for text in file.paragraphs:
            print(text.text)
    except:
        pass

path = "D:\\OUT" #文件夹目录
files= os.listdir(path) #得到文件夹下的所有文件名称
s ={}
for file in files: #遍历文件夹
     if not os.path.isdir(file): #判断是否是文件夹，不是文件夹才打开
         name = file.split('.')[0].replace('~','').replace('$','').replace('答案','')
         s[name]=''
print(s)


for key in s:
    getText(key)
